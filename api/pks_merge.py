"""
PKS Mail Merge — Vercel Python Serverless Function
===================================================

Vercel auto-detects Python files in /api folder dan deploy sebagai
serverless function terpisah dari Next.js.

Endpoint: POST /api/pks_merge
Body: { "data": {placeholder_name: value, ...} }
Response: binary .docx file (download)

Vercel Python function format: pakai handler class BaseHTTPRequestHandler
atau Flask. Saya pakai Flask-style dengan Vercel's Python runtime.
"""

import json
import io
import os
import re
import sys
import traceback
from pathlib import Path

# Add parent to path untuk import pks_merge
sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

# ============================================================
# Mail merge logic (inline — Vercel Python function tidak bisa
# import dari /scripts dengan mudah, jadi inline saja)
# ============================================================

PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
EMPTY_PLACEHOLDER = "___"


def get_replacement_value(token, data):
    if token in data:
        value = data[token]
        if value is None:
            return EMPTY_PLACEHOLDER
        value = str(value)
        if value.strip() == "":
            return EMPTY_PLACEHOLDER
        return value
    return None


def replace_in_paragraph(paragraph, data, stats):
    for run in paragraph.runs:
        matches = list(PATTERN.finditer(run.text))
        if not matches:
            continue
        new_text = run.text
        for m in reversed(matches):
            token = m.group(1)
            value = get_replacement_value(token, data)
            if value is not None:
                new_text = new_text[:m.start()] + value + new_text[m.end():]
                stats["replaced"] += 1
                if value == EMPTY_PLACEHOLDER:
                    stats["empty_filled"] += 1
            else:
                stats["missing_data"] += 1
                stats["missing_tokens"].add(token)
        run.text = new_text


def replace_in_cell(cell, data, stats):
    for p in cell.paragraphs:
        replace_in_paragraph(p, data, stats)
    for nt in cell.tables:
        for row in nt.rows:
            for c in row.cells:
                replace_in_cell(c, data, stats)


def replace_in_header_footer(hdr_ftr, data, stats):
    for p in hdr_ftr.paragraphs:
        replace_in_paragraph(p, data, stats)
    for tbl in hdr_ftr.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)


def validate_structure_unchanged(original_doc, merged_doc):
    def count_structure(doc):
        body = doc.element.body
        # python-docx 1.2.0 pakai lxml
        from lxml import etree
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        return {
            "body_paragraphs": len(body.findall('w:p', ns)),
            "body_tables": len(body.findall('w:tbl', ns)),
            "all_paragraphs": len(list(body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'))),
            "all_tables": len(list(body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'))),
            "all_runs": len(list(body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'))),
        }
    orig = count_structure(original_doc)
    merged = count_structure(merged_doc)
    diff = {k: merged[k] - orig[k] for k in orig}
    is_valid = all(v == 0 for v in diff.values())
    return is_valid, diff


def perform_mail_merge(template_path, data):
    from docx import Document

    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    original_doc = Document(template_path)
    doc = Document(template_path)

    stats = {
        "replaced": 0,
        "empty_filled": 0,
        "missing_data": 0,
        "missing_tokens": set(),
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, data, stats)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            replace_in_header_footer(hdr_ftr, data, stats)

    is_valid, diff = validate_structure_unchanged(original_doc, doc)
    if not is_valid:
        raise ValueError(f"Structure changed during merge! Diff: {diff}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    stats["missing_tokens"] = sorted(stats["missing_tokens"])
    stats["structure_valid"] = is_valid
    stats["structure_diff"] = diff

    return output.getvalue(), stats


# ============================================================
# Vercel Python serverless function entry point
# ============================================================
# Vercel Python function format: harus ada `handler` class yang extend
# BaseHTTPRequestHandler, ATAU pakai framework (Flask/FastAPI) yang
# Vercel auto-detect.
#
# Format paling sederhana: pakai http.server.BaseHTTPRequestHandler

from http.server import BaseHTTPRequestHandler
from urllib.parse import urlparse


class handler(BaseHTTPRequestHandler):
    """Vercel Python serverless function handler."""

    def do_GET(self):
        """Health check + list placeholders."""
        parsed = urlparse(self.path)
        if parsed.path == '/api/pks_merge':
            # List placeholders
            try:
                template_path = self._get_template_path()
                if not template_path or not Path(template_path).exists():
                    self._send_json(404, {"error": "Template not found", "path": template_path})
                    return
                from docx import Document
                doc = Document(template_path)
                placeholders = set()
                for p in doc.paragraphs:
                    placeholders.update(PATTERN.findall(p.text))
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                placeholders.update(PATTERN.findall(p.text))
                for section in doc.sections:
                    for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                                    section.footer, section.first_page_footer, section.even_page_footer]:
                        for p in hdr_ftr.paragraphs:
                            placeholders.update(PATTERN.findall(p.text))
                self._send_json(200, {
                    "ok": True,
                    "service": "pks_merge_python",
                    "placeholders": sorted(placeholders),
                    "count": len(placeholders),
                })
            except Exception as e:
                self._send_json(500, {"error": str(e), "trace": traceback.format_exc()})
            return

        self._send_json(404, {"error": "Not found"})

    def do_POST(self):
        """Perform mail merge — return binary .docx atau JSON dengan stats."""
        try:
            # Parse body
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length).decode('utf-8')
            payload = json.loads(body)

            data = payload.get('data', {})
            return_stats = payload.get('return_stats', False)
            template_path = self._get_template_path()

            if not template_path or not Path(template_path).exists():
                self._send_json(500, {
                    "error": "Template not found",
                    "path": template_path,
                    "hint": "Pastikan templates/pks_template_bersih.docx ada di repo"
                })
                return

            # Perform merge
            merged_bytes, stats = perform_mail_merge(template_path, data)

            if return_stats:
                import base64
                self._send_json(200, {
                    "ok": True,
                    "stats": stats,
                    "file_base64": base64.b64encode(merged_bytes).decode('utf-8'),
                    "filename": f"PKS_{(data.get('NAMA_FASKES') or 'draft').replace(' ', '_')}.docx",
                    "bytes": len(merged_bytes),
                })
                return

            # Return binary .docx
            filename = f"PKS_{(data.get('NAMA_FASKES') or 'draft').replace(' ', '_')}.docx"
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.send_header('Content-Length', str(len(merged_bytes)))
            self.send_header('X-Stats-Replaced', str(stats['replaced']))
            self.send_header('X-Stats-Empty-Filled', str(stats['empty_filled']))
            self.send_header('X-Stats-Missing', str(stats['missing_data']))
            self.send_header('X-Stats-Structure-Valid', str(stats['structure_valid']))
            self.end_headers()
            self.wfile.write(merged_bytes)

        except Exception as e:
            self._send_json(500, {
                "error": str(e),
                "trace": traceback.format_exc().split('\n')[-5:]
            })

    def _get_template_path(self):
        """Cari template file di beberapa lokasi yang mungkin."""
        # Vercel Python function working dir biasanya /var/task/
        candidates = [
            '/var/task/templates/pks_template_bersih.docx',
            '/tmp/templates/pks_template_bersih.docx',
            os.path.join(os.path.dirname(__file__), '..', 'templates', 'pks_template_bersih.docx'),
            os.path.join(os.getcwd(), 'templates', 'pks_template_bersih.docx'),
        ]
        for c in candidates:
            if Path(c).exists():
                return c
        return None

    def _send_json(self, status, data):
        body = json.dumps(data, indent=2, ensure_ascii=False).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)
