"""
Mail Merge Script untuk PKS Template (versi bersih) — v2
========================================================

Perubahan dari v1:
1. Data values diperbaiki:
   - NAMA_KANTOR_CABANG: "Cirebon" → "CIREBON" (all-caps, konsisten dengan "KANTOR CABANG")
   - JABATAN_PENANDATANGAN_PIHAK_KEDUA: typo "Djuaansih" → "Djuansih"
   - NAMA_SAKSI_PIHAK_KEDUA: "Tiar Salman, S.T, M.M" → "Tiar Salman, S.T., M.M." (titik after gelar)
   - NAMA_PIC_ADMIN_FASKES: "lin Siti Hindasah" → "Lin Siti Hindasah" (kapital L)
2. Empty value handling: jika value kosong, ganti placeholder dengan "______"
   (garis bawah, standar form Indonesia untuk diisi manual)
3. Better logging: tampilkan berapa placeholder terisi vs kosong
"""

from docx import Document
from pathlib import Path
import re
import sys

TEMPLATE_PATH = "/home/z/my-project/download/pks_template_bersih.docx"
OUTPUT_PATH = "/home/z/my-project/download/pks_hasil_mailmerge.docx"

# Konstanta untuk empty value
EMPTY_PLACEHOLDER = "___"  # 3 underscores — pendek supaya tidak wrap ke baris baru

# ============= DATA UNTUK MAIL MERGE =============
data = {
    # ----- PIHAK PERTAMA (BPJS Ketenagakerjaan) -----
    "NAMA_KANTOR_CABANG": "CIREBON",  # all-caps, konsisten dengan "KANTOR CABANG"
    "ALAMAT_KANTOR_CABANG": "Jalan Evakuasi Nomor 11B Karyamulya Kota Cirebon",
    "TELP_FAX_BPJS": "(0231) 485660",
    "NAMA_KEPALA_KANTOR_CABANG": "Ahmad Feisal Santoso",
    "NOMOR_KEP_DIREKSI": "KEP/11/012025",
    "JUDUL_KEP_DIREKSI": "Mutasi Pejabat",
    "NOMOR_SURAT_KUASA": "SKS/04/012026",
    "TANGGAL_SURAT_KUASA": "21 Januari 2025",
    "JABATAN_PENANDATANGAN_PIHAK_PERTAMA": "KEPALA KANTOR CABANG",
    "NOMOR_PKS_PIHAK_PERTAMA": "PER/11/032026",
    "NOMOR_PKS_SEBELUMNYA_PIHAK_PERTAMA": "PER/09/032025",
    "NAMA_SAKSI_PIHAK_PERTAMA": "Rohaya Hartisari",
    "JABATAN_SAKSI_PIHAK_PERTAMA": "Manajer Kasus KK-PAK",
    "NAMA_PIC_BPJS": "Rohaya Hartisari",
    "JABATAN_PIC_BPJS": "Manajer Kasus KK-PAK",
    "HP_PIC_BPJS": "081212331065",
    "EMAIL_PIC_BPJS": "laporkk03.neo@gmail.com",

    # ----- PIHAK KEDUA (Faskes) -----
    "NAMA_FASKES": "Rumah Sakit Djuansih Majalengka",
    "BENTUK_FASKES": "Fasilitas Pelayanan Kesehatan",
    "JENIS_FASKES": "RS Umum Swasta Tipe C",
    "ALAMAT_FASKES": "Jalan Siliwangi Kilometer 7 Nomor 84 Desa Karyamukti Kecamatan Panyingkiran Kabupaten Majalengka Provinsi Jawa Barat, 45459",
    "WEB_FASKES": "www.rsdjuansih.com",
    "TELP_FAX_FASKES": "(0233) 3603030",
    "JENIS_AKTA_PENDIRIAN": "Akta Pendirian",
    "NOMOR_AKTA_PENDIRIAN": "05",
    "TANGGAL_AKTA_PENDIRIAN": "16 Juni 2026",
    "DASAR_KEWENANGAN_PIHAK_KEDUA": "Surat Keputusan Direktur PT. Santosha Adyatama Husada Nomor: 21C/DIR/PTSAH/V/2025 tanggal 21 Mei 2025",
    "NAMA_PENANDATANGAN_PIHAK_KEDUA": "dr. Oky Trisdiana Wahyat, Sp.B.,MMRS.,CHMP",
    "JABATAN_PENANDATANGAN_PIHAK_KEDUA": "Direktur Rumah Sakit Djuansih Majalengka",  # fixed typo Djuaansih → Djuansih
    "NOMOR_PKS_PIHAK_KEDUA": "002/PJN/RSDM/11/2026",
    "NOMOR_PKS_SEBELUMNYA_PIHAK_KEDUA": "002/PJN/RSDM/09/2025",
    "PERIHAL_PKS_SEBELUMNYA": "Pelaksanaan Pelayanan Kesehatan bagi Peserta Program Jaminan Kecelakaan Kerja",
    "TANGGAL_BERAKHIR_PKS_SEBELUMNYA": "10 April 2027",
    "NAMA_SAKSI_PIHAK_KEDUA": "Tiar Salman, S.T., M.M.",  # fixed: tambah titik after S.T dan M.M
    "JABATAN_SAKSI_PIHAK_KEDUA": "Manajer Bisnis",
    "NAMA_PIC_ADMIN_FASKES": "Lin Siti Hindasah, SE",  # fixed: kapital L
    "JABATAN_PIC_ADMIN_FASKES": "Manajer Ketatausahaan",
    "HP_PIC_ADMIN_FASKES": "082316668554",
    "EMAIL_PIC_ADMIN_FASKES": "linsitihindasah123@gmail.com",
    "NAMA_PIC_KLINIS_FASKES": "dr. Irma Gianova Lestari",
    "JABATAN_PIC_KLINIS_FASKES": "Manajer Pelayanan Medis",
    "HP_PIC_KLINIS_FASKES": "081271172013",
    "EMAIL_PIC_KLINIS_FASKES": "drirma.gianova@gmail.com",

    # ----- Tarif & Pelayanan -----
    "JENIS_TARIF_KK_PAK": "Tarif fee for services",
    "NAMA_RS_PEMERINTAH_ACUAN": "Rumah Sakit Umum Daerah Majalengka",
    "KELAS_RAWAT_INAP_KK_PAK": "2",
    "NOMOR_BA_NEGOSIASI": "BA/16/032026",

    # ----- Tanggal Tanda Tangan -----
    "HARI_TANDA_TANGAN": "Rabu",
    "TANGGAL_TANDA_TANGAN": "1",
    "BULAN_TANDA_TANGAN": "April",
    "TAHUN_TANDA_TANGAN": "2026",
    "KOTA_TANDA_TANGAN": "Majalengka",
    "KOTA_PENGADILAN_NEGERI": "Cirebon",
    "TANGGAL_MULAI_PKS": "10 April 2026",
    "TANGGAL_BERAKHIR_PKS": "10 April 2027",

    # ----- BA Negosiasi -----
    "HARI_NEGOSIASI": "Kamis",
    "TANGGAL_NEGOSIASI": "5",
    "BULAN_NEGOSIASI": "Maret",
    "TAHUN_NEGOSIASI": "2026",
    "JAM_NEGOSIASI": "10:10",
    "TAHUN_TARIF_NEGOSIASI": "2026",
    "TANGGAL_PENAWARAN": "5",
    "BULAN_PENAWARAN": "November",
    "TAHUN_PENAWARAN": "2025",

    # ----- BA Rekonsiliasi (diisi saat event rekonsiliasi terjadi) -----
    # Dibiarkan kosong — saat mail merge akan diganti dengan "______" untuk diisi manual
    "NOMOR_BA_REKONSILIASI": "",
    "TANGGAL_REKONSILIASI": "",
    "BULAN_REKONSILIASI": "",
    "TAHUN_REKONSILIASI": "",
    "BULAN_AWAL_REKONSILIASI": "",
    "BULAN_AKHIR_REKONSILIASI": "",

    # ----- Informasi Kelengkapan Dokumen (diisi saat event tagihan terjadi) -----
    "NOMOR_INFORMASI_KELENGKAPAN": "",
    "TANGGAL_INFORMASI_KELENGKAPAN": "",
    "BULAN_INFORMASI_KELENGKAPAN": "",
    "TAHUN_INFORMASI_KELENGKAPAN": "",
    "BULAN_PELAYANAN": "",
    "TAHUN_PELAYANAN": "",
    "JUMLAH_KASUS_TIDAK_LENGKAP": "",
    "BATAS_HARI_PENLENGKAPAN": "",

    # ----- Pakta Integritas -----
    "TEMPAT_PAKTA": "Cirebon",
    "BULAN_PAKTA": "Maret",
    "TAHUN_PAKTA": "2026",
    "NAMA_PIMPINAN_FASKES": "dr. Oky Trisdiana Wahyat, Sp.B.,MMRS.,CHMP",
    "JABATAN_PIMPINAN_FASKES": "Direktur",
    "NAMA_PIC_USER_EPLKK": "Tiar Salman, S.T., M.M.",
    "JABATAN_PIC_USER_EPLKK": "Manajer Bisnis",
    "NAMA_PIC_NARAHUBUNG": "Lin Siti Hindasah, SE",
    "JABATAN_PIC_NARAHUBUNG": "Manajer Ketatausahaan",

    # ----- Rekening Bank -----
    "NAMA_REKENING": "RS Djuansih Majalengka",
    "NOMOR_REKENING": "3333668881",
    "NAMA_BANK": "Bank Syariah Indonesia (BSI)",
    "CABANG_BANK": "Kadipaten - Majalengka",
}

# ============= LOGIC =============

PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

def get_replacement_value(token, data):
    """Get the value to replace placeholder with.
    If token exists in data and value is non-empty, return value.
    If token exists but value is empty, return EMPTY_PLACEHOLDER (______).
    If token doesn't exist in data, return None (leave placeholder)."""
    if token in data:
        value = str(data[token])
        if value.strip() == "":
            return EMPTY_PLACEHOLDER
        return value
    return None

def replace_in_paragraph(paragraph, data, stats):
    """Replace placeholders in a single paragraph. Update stats counter."""
    for run in paragraph.runs:
        # Find all placeholders in this run
        matches = list(PATTERN.finditer(run.text))
        if not matches:
            continue
        # Replace from end to start to preserve indices
        new_text = run.text
        for m in reversed(matches):
            token = m.group(1)
            placeholder = m.group(0)
            value = get_replacement_value(token, data)
            if value is not None:
                new_text = new_text[:m.start()] + value + new_text[m.end():]
                stats["replaced"] += 1
                if value == EMPTY_PLACEHOLDER:
                    stats["empty_filled"] += 1
            else:
                stats["missing_data"] += 1
                stats["missing_tokens"].add(token)
        run.text = new_text

def replace_in_cell(cell, data, stats):
    """Replace in a cell + recurse into nested tables."""
    for p in cell.paragraphs:
        replace_in_paragraph(p, data, stats)
    for nt in cell.tables:
        for row in nt.rows:
            for c in row.cells:
                replace_in_cell(c, data, stats)

def replace_in_header_footer(hdr_ftr, data, stats):
    for p in hdr_ftr.paragraphs:
        replace_in_paragraph(p, data, stats)
    for tbl in hdr_ftr.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

def main():
    if not Path(TEMPLATE_PATH).exists():
        print(f"❌ Template tidak ditemukan: {TEMPLATE_PATH}")
        sys.exit(1)

    print(f"📖 Loading template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    stats = {
        "replaced": 0,
        "empty_filled": 0,
        "missing_data": 0,
        "missing_tokens": set(),
    }

    # Body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, data, stats)

    # All tables (with nested)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

    # Headers & footers
    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            replace_in_header_footer(hdr_ftr, data, stats)

    print(f"\n📊 Statistik Mail Merge:")
    print(f"   ✅ Placeholder terisi data: {stats['replaced'] - stats['empty_filled']}")
    print(f"   📝 Placeholder kosong → '______' (untuk diisi manual): {stats['empty_filled']}")
    print(f"   ⚠️  Placeholder tidak punya data key: {stats['missing_data']}")
    if stats["missing_tokens"]:
        print(f"   Tokens tanpa data key:")
        for t in sorted(stats["missing_tokens"]):
            print(f"     - {{{{{t}}}}}")

    print(f"\n💾 Saving: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print(f"✅ Done!")

if __name__ == "__main__":
    main()
