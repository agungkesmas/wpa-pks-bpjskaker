#!/usr/bin/env python3
"""
PKS Merge Script — CLI / stdin-stdout mode
==========================================

Script ini cocok untuk dipanggil dari:
- Next.js API route (via child_process.spawn)
- Bash script
- Cron job
- Web backend (Node.js, PHP, Ruby, dll)

Cara pakai:
    cat data.json | python3 pks_merge.py --template templates/pks_template_bersih.docx > result.docx

Atau:

    python3 pks_merge.py --template templates/pks_template_bersih.docx --input data.json --output result.docx

Input: JSON berisi dictionary {placeholder_name: value}
Output: binary .docx ke stdout (atau file jika --output diset)
"""

import sys
import json
import argparse
import io
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# ============================================================
# Core mail merge logic — identical to pks_merge_api.py
# ============================================================

PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
EMPTY_PLACEHOLDER = "___"  # 3 underscores — pendek, tidak bikin Word wrap

def get_replacement_value(token, data):
    if token in data:
        value = str(data[token]) if data[token] is not None else ""
        if value.strip() == "":
            return EMPTY_PLACEHOLDER
        return value
    return None

def replace_in_paragraph(paragraph, data, stats):
    for run in paragraph.runs:
        matches = list(PATTERN.finditer(run.text))
        if not matches:
            continue
        new_text = run.text
        for m in reversed(matches):
            token = m.group(1)
            placeholder = m.group(0)
            value = get_replacement_value(token, data)
            if value is not None:
                new_text = new_text[:m.start()] + value + new_text[m.end():]
                stats["replaced"] += 1
                if value == EMPTY_PLACEHOLDER:
                    stats["empty_filled"] += 1
            else:
                stats["missing_data"] += 1
                stats["missing_tokens"].add(token)
        run.text = new_text

def replace_in_cell(cell, data, stats):
    for p in cell.paragraphs:
        replace_in_paragraph(p, data, stats)
    for nt in cell.tables:
        for row in nt.rows:
            for c in row.cells:
                replace_in_cell(c, data, stats)

def replace_in_header_footer(hdr_ftr, data, stats):
    for p in hdr_ftr.paragraphs:
        replace_in_paragraph(p, data, stats)
    for tbl in hdr_ftr.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

def validate_structure_unchanged(original_doc, merged_doc):
    def count_structure(doc):
        body = doc.element.body
        return {
            "body_paragraphs": sum(1 for c in body.iterchildren() if c.tag == qn('w:p')),
            "body_tables": sum(1 for c in body.iterchildren() if c.tag == qn('w:tbl')),
            "all_paragraphs": len(list(body.iter(qn('w:p')))),
            "all_tables": len(list(body.iter(qn('w:tbl')))),
            "all_runs": len(list(body.iter(qn('w:r')))),
        }
    orig = count_structure(original_doc)
    merged = count_structure(merged_doc)
    diff = {k: merged[k] - orig[k] for k in orig}
    is_valid = all(v == 0 for v in diff.values())
    return is_valid, diff

def perform_mail_merge(template_path, data):
    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    original_doc = Document(template_path)
    doc = Document(template_path)
    
    stats = {
        "replaced": 0,
        "empty_filled": 0,
        "missing_data": 0,
        "missing_tokens": set(),
    }
    
    for p in doc.paragraphs:
        replace_in_paragraph(p, data, stats)
    
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)
    
    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            replace_in_header_footer(hdr_ftr, data, stats)
    
    is_valid, diff = validate_structure_unchanged(original_doc, doc)
    if not is_valid:
        raise ValueError(f"Structure changed during merge! Diff: {diff}")
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    stats["missing_tokens"] = sorted(stats["missing_tokens"])
    stats["structure_valid"] = is_valid
    stats["structure_diff"] = diff
    
    return output.getvalue(), stats

# ============================================================
# CLI entry point
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="PKS Mail Merge — CLI mode")
    parser.add_argument("--template", required=True, help="Path to .docx template")
    parser.add_argument("--input", help="Path to input JSON file (if not using stdin)")
    parser.add_argument("--output", help="Path to output .docx file (if not using stdout)")
    parser.add_argument("--stats", help="Path to write stats JSON (optional)")
    args = parser.parse_args()
    
    # Read input data
    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        # Read from stdin
        stdin_text = sys.stdin.read()
        if not stdin_text.strip():
            print("Error: No input provided. Use --input or pipe JSON via stdin.", file=sys.stderr)
            sys.exit(1)
        data = json.loads(stdin_text)
    
    # Perform merge
    try:
        merged_bytes, stats = perform_mail_merge(args.template, data)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    
    # Write output
    if args.output:
        with open(args.output, "wb") as f:
            f.write(merged_bytes)
        print(f"✅ Saved: {args.output}", file=sys.stderr)
    else:
        # Write to stdout (binary)
        sys.stdout.buffer.write(merged_bytes)
    
    # Write stats if requested
    if args.stats:
        with open(args.stats, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
    
    # Print stats to stderr (so it doesn't corrupt stdout binary)
    print(f"📊 Stats:", file=sys.stderr)
    print(f"   Replaced: {stats['replaced'] - stats['empty_filled']}", file=sys.stderr)
    print(f"   Empty filled with '___': {stats['empty_filled']}", file=sys.stderr)
    print(f"   Missing data: {stats['missing_data']}", file=sys.stderr)
    print(f"   Structure valid: {stats['structure_valid']}", file=sys.stderr)

if __name__ == "__main__":
    main()
