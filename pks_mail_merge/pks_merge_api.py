"""
PKS Mail Merge API — Backend Service
=====================================

File ini berisi 3 varian integrasi yang bisa dipilih:
1. FastAPI (modern, async, recommended untuk production)
2. Flask (sederhana, legacy-friendly)
3. Next.js API Route (TypeScript — untuk integrasi langsung dengan Next.js)

Pilih salah satu sesuai stack web Anda.

PRINSIP UTAMA — "Tidak Bergeser Satu Milimeter Pun":
=====================================================
1. HANYA mengganti teks di dalam <w:t> element (text runs)
2. TIDAK menambah/menghapus paragraph, run, atau table
3. TIDAK mengubah style, font, alignment, spacing
4. Validasi struktur XML sebelum & sesudah — jumlah paragraph, table, run harus identik
5. Jika ada placeholder kosong (empty value), ganti dengan "___" (3 underscore, jangan 6+ agar tidak wrap)

Cara pakai:
- Simpan template di: ./templates/pks_template_bersih.docx
- Kirim POST request dengan JSON berisi data field
- Response: file .docx hasil mail merge
"""

import re
import io
import json
import hashlib
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.oxml.ns import qn

# ============================================================
# CORE: Mail Merge Function (dipakai semua varian)
# ============================================================

PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
EMPTY_PLACEHOLDER = "___"  # 3 underscores — pendek, tidak bikin Word wrap

# Path template — relatif terhadap root project
TEMPLATE_PATH = Path("templates/pks_template_bersih.docx")

def get_replacement_value(token: str, data: Dict[str, Any]) -> str:
    """Get value to replace placeholder with.
    - Jika token ada di data dan value non-empty: return value (str)
    - Jika token ada tapi value empty: return EMPTY_PLACEHOLDER ("___")
    - Jika token tidak ada di data: return None (placeholder dibiarkan)
    """
    if token in data:
        value = str(data[token]) if data[token] is not None else ""
        if value.strip() == "":
            return EMPTY_PLACEHOLDER
        return value
    return None

def replace_in_paragraph(paragraph, data: Dict[str, Any], stats: Dict[str, int]):
    """Replace placeholders in a single paragraph. Update stats counter."""
    for run in paragraph.runs:
        matches = list(PATTERN.finditer(run.text))
        if not matches:
            continue
        # Replace from end to start to preserve indices
        new_text = run.text
        for m in reversed(matches):
            token = m.group(1)
            placeholder = m.group(0)
            value = get_replacement_value(token, data)
            if value is not None:
                new_text = new_text[:m.start()] + value + new_text[m.end():]
                stats["replaced"] += 1
                if value == EMPTY_PLACEHOLDER:
                    stats["empty_filled"] += 1
            else:
                stats["missing_data"] += 1
                stats["missing_tokens"].add(token)
        run.text = new_text

def replace_in_cell(cell, data: Dict[str, Any], stats: Dict[str, int]):
    """Replace in a cell + recurse into nested tables."""
    for p in cell.paragraphs:
        replace_in_paragraph(p, data, stats)
    for nt in cell.tables:
        for row in nt.rows:
            for c in row.cells:
                replace_in_cell(c, data, stats)

def replace_in_header_footer(hdr_ftr, data: Dict[str, Any], stats: Dict[str, int]):
    for p in hdr_ftr.paragraphs:
        replace_in_paragraph(p, data, stats)
    for tbl in hdr_ftr.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

def validate_structure_unchanged(original_doc, merged_doc):
    """Validate that XML structure is identical between original and merged document.
    Returns (is_valid, diff_dict)."""
    def count_structure(doc):
        body = doc.element.body
        return {
            "body_paragraphs": sum(1 for c in body.iterchildren() if c.tag == qn('w:p')),
            "body_tables": sum(1 for c in body.iterchildren() if c.tag == qn('w:tbl')),
            "all_paragraphs": len(list(body.iter(qn('w:p')))),
            "all_tables": len(list(body.iter(qn('w:tbl')))),
            "all_runs": len(list(body.iter(qn('w:r')))),
        }
    orig = count_structure(original_doc)
    merged = count_structure(merged_doc)
    diff = {k: merged[k] - orig[k] for k in orig}
    is_valid = all(v == 0 for v in diff.values())
    return is_valid, diff

def perform_mail_merge(template_path: Path, data: Dict[str, Any]) -> tuple[bytes, Dict[str, Any]]:
    """
    Perform mail merge on a PKS template.
    
    Args:
        template_path: Path to .docx template file
        data: Dictionary of {placeholder_name: value}
    
    Returns:
        Tuple of (merged_docx_bytes, stats_dict)
        stats_dict contains: replaced, empty_filled, missing_data, missing_tokens, structure_valid, structure_diff
    
    Raises:
        FileNotFoundError: if template not found
        ValueError: if structure changed (should never happen with our approach)
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Load original template for structure validation
    original_doc = Document(template_path)
    
    # Load working copy
    doc = Document(template_path)
    
    stats = {
        "replaced": 0,
        "empty_filled": 0,
        "missing_data": 0,
        "missing_tokens": set(),
    }
    
    # Body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, data, stats)
    
    # All tables (with nested)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)
    
    # Headers & footers
    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            replace_in_header_footer(hdr_ftr, data, stats)
    
    # Validate structure unchanged
    is_valid, diff = validate_structure_unchanged(original_doc, doc)
    if not is_valid:
        raise ValueError(f"Structure changed during merge! Diff: {diff}")
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Convert set to list for JSON serialization
    stats["missing_tokens"] = sorted(stats["missing_tokens"])
    stats["structure_valid"] = is_valid
    stats["structure_diff"] = diff
    
    return output.getvalue(), stats


# ============================================================
# VARIAN 1: FastAPI (recommended untuk production)
# ============================================================
"""
Cara pakai:
1. pip install fastapi uvicorn python-docx
2. Simpan file ini sebagai app.py
3. Buat folder templates/, simpan pks_template_bersih.docx di dalamnya
4. Jalankan: uvicorn app:app --reload
5. Test: curl -X POST http://localhost:8000/merge -H "Content-Type: application/json" -d @data.json --output result.docx
"""

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import os

app = FastAPI(title="PKS Mail Merge API", version="1.0.0")

# CORS — sesuaikan origin yang diizinkan
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Ganti dengan domain Anda di production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class MergeRequest(BaseModel):
    """Schema untuk request body mail merge."""
    data: Dict[str, Any]
    template_name: Optional[str] = "pks_template_bersih.docx"

@app.post("/merge")
async def merge_pks(request: MergeRequest):
    """
    Endpoint untuk mail merge PKS.
    Kirim JSON dengan field "data" berisi dictionary placeholder → value.
    """
    try:
        template_path = Path("templates") / request.template_name
        merged_bytes, stats = perform_mail_merge(template_path, request.data)
        
        # Generate filename
        faskes_name = request.data.get("NAMA_FASKES", "pks").replace(" ", "_")
        filename = f"PKS_{faskes_name}.docx"
        
        return StreamingResponse(
            io.BytesIO(merged_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            background=None,
        )
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=500, detail=f"Structure validation failed: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/merge-with-stats")
async def merge_pks_with_stats(request: MergeRequest):
    """
    Endpoint yang mengembalikan file + statistik sebagai JSON.
    File di-encode base64 dalam response.
    """
    import base64
    try:
        template_path = Path("templates") / request.template_name
        merged_bytes, stats = perform_mail_merge(template_path, request.data)
        
        return JSONResponse({
            "success": True,
            "stats": stats,
            "file_base64": base64.b64encode(merged_bytes).decode("utf-8"),
            "filename": f"PKS_{request.data.get('NAMA_FASKES', 'pks').replace(' ', '_')}.docx",
        })
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)

@app.get("/template-fields")
async def get_template_fields():
    """Return list of all placeholders in the template."""
    try:
        template_path = Path("templates") / "pks_template_bersih.docx"
        doc = Document(template_path)
        placeholders = set()
        
        # Scan all paragraphs
        for p in doc.paragraphs:
            for m in PATTERN.finditer(p.text):
                placeholders.add(m.group(1))
        
        # Scan all tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for m in PATTERN.finditer(p.text):
                            placeholders.add(m.group(1))
                    # Nested tables
                    for nt in cell.tables:
                        for nrow in nt.rows:
                            for ncell in nrow.cells:
                                for p in ncell.paragraphs:
                                    for m in PATTERN.finditer(p.text):
                                        placeholders.add(m.group(1))
        
        return {"placeholders": sorted(placeholders), "count": len(placeholders)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health():
    return {"status": "ok", "template_exists": TEMPLATE_PATH.exists()}


# ============================================================
# VARIAN 2: Flask (sederhana, legacy-friendly)
# ============================================================
"""
Cara pakai:
1. pip install flask python-docx
2. Simpan file ini sebagai app.py
3. Buat folder templates/, simpan pks_template_bersih.docx di dalamnya
4. Jalankan: python app.py
5. Test: curl -X POST http://localhost:5000/merge -H "Content-Type: application/json" -d @data.json --output result.docx
"""

# Uncomment baris di bawah untuk pakai Flask:
# from flask import Flask, request, jsonify, send_file
# flask_app = Flask(__name__)
# 
# @flask_app.route("/merge", methods=["POST"])
# def merge_pks_flask():
#     try:
#         data = request.json.get("data", {})
#         template_name = request.json.get("template_name", "pks_template_bersih.docx")
#         template_path = Path("templates") / template_name
#         merged_bytes, stats = perform_mail_merge(template_path, data)
#         
#         faskes_name = data.get("NAMA_FASKES", "pks").replace(" ", "_")
#         filename = f"PKS_{faskes_name}.docx"
#         
#         return send_file(
#             io.BytesIO(merged_bytes),
#             as_attachment=True,
#             download_name=filename,
#             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         )
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500
# 
# if __name__ == "__main__":
#     flask_app.run(host="0.0.0.0", port=5000, debug=True)


# ============================================================
# VARIAN 3: Next.js API Route (TypeScript)
# ============================================================
"""
Untuk Next.js, simpan kode berikut di: app/api/pks-merge/route.ts

```typescript
import { NextRequest, NextResponse } from 'next/server';
import { spawn } from 'child_process';
import path from 'path';
import fs from 'fs/promises';

export async function POST(request: NextRequest) {
  try {
    const body = await request.json();
    const { data, templateName = 'pks_template_bersih.docx' } = body;
    
    // Panggil Python script (lebih aman & teruji)
    const result = await new Promise<Buffer>((resolve, reject) => {
      const python = spawn('python3', [
        path.join(process.cwd(), 'scripts', 'pks_merge.py'),
        '--template', path.join(process.cwd(), 'templates', templateName),
      ]);
      
      const chunks: Buffer[] = [];
      let stderr = '';
      
      python.stdin.write(JSON.stringify(data));
      python.stdin.end();
      
      python.stdout.on('data', (chunk) => chunks.push(Buffer.from(chunk)));
      python.stderr.on('data', (chunk) => stderr += chunk.toString());
      python.on('close', (code) => {
        if (code !== 0) reject(new Error(stderr));
        else resolve(Buffer.concat(chunks));
      });
    });
    
    const faskesName = (data.NAMA_FASKES || 'pks').replace(/\s+/g, '_');
    const filename = `PKS_${faskesName}.docx`;
    
    return new NextResponse(result, {
      headers: {
        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'Content-Disposition': `attachment; filename="${filename}"`,
      },
    });
  } catch (error) {
    return NextResponse.json({ error: String(error) }, { status: 500 });
  }
}
```

Lalu buat script Python `scripts/pks_merge.py` yang membaca JSON dari stdin dan menulis docx ke stdout.
"""


# ============================================================
# CLI mode — untuk testing tanpa web server
# ============================================================
"""
Cara pakai via CLI:
1. Simpan data di file JSON (mis. data.json)
2. Jalankan: python app.py --cli data.json output.docx
"""
if __name__ == "__main__":
    import sys
    import argparse
    
    parser = argparse.ArgumentParser(description="PKS Mail Merge")
    parser.add_argument("--cli", action="store_true", help="CLI mode")
    parser.add_argument("input_json", help="Path to JSON data file")
    parser.add_argument("output_docx", help="Path to output .docx file")
    parser.add_argument("--template", default="templates/pks_template_bersih.docx", help="Template path")
    args = parser.parse_args()
    
    if args.cli:
        with open(args.input_json, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        template_path = Path(args.template)
        merged_bytes, stats = perform_mail_merge(template_path, data)
        
        with open(args.output_docx, "wb") as f:
            f.write(merged_bytes)
        
        print(f"✅ Saved: {args.output_docx}")
        print(f"📊 Stats:")
        print(f"   Replaced: {stats['replaced'] - stats['empty_filled']}")
        print(f"   Empty filled with '___': {stats['empty_filled']}")
        print(f"   Missing data: {stats['missing_data']}")
        if stats['missing_tokens']:
            print(f"   Missing tokens: {stats['missing_tokens']}")
        print(f"   Structure valid: {stats['structure_valid']}")
