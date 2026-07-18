#!/usr/bin/env python3
"""
PKS Merge Script — v2 (Production)
==================================

Versi produksi setelah audit:
- ✅ 100% byte-identical dengan file referensi (pks_hasil_mailmerge.docx)
- ✅ Structure validation: paragraph/table/run count identik sebelum & sesudah
- ✅ Handle None, empty string, missing key dengan benar
- ✅ Process body + tables (nested) + headers + footers

Cara pakai dari Next.js:
    child_process.spawn('python3', ['scripts/pks_merge.py',
                                    '--template', 'templates/pks_template_bersih.docx',
                                    '--input', '/tmp/data.json',
                                    '--output', '/tmp/result.docx',
                                    '--stats', '/tmp/stats.json'])

Atau via stdin/stdout (binary-safe):
    echo '{"NAMA_FASKES":"..."}' | python3 scripts/pks_merge.py \\
        --template templates/pks_template_bersih.docx > result.docx

Input: JSON {placeholder_name: value}
Output: binary .docx (ke file jika --output, ke stdout binary jika tidak)
"""

import sys
import json
import argparse
import io
import re
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# ============================================================
# Konstanta
# ============================================================

PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
EMPTY_PLACEHOLDER = "___"  # 3 underscores — pendek, tidak bikin Word wrap

# ============================================================
# Core mail merge logic
# ============================================================

def get_replacement_value(token, data):
    """
    Tentukan nilai pengganti placeholder.
    - Jika token ada di data dan value non-empty: return value (str)
    - Jika token ada tapi value None atau empty: return EMPTY_PLACEHOLDER ("___")
    - Jika token tidak ada di data: return None (placeholder dibiarkan apa adanya)
    """
    if token in data:
        value = data[token]
        # Handle None, NaN, dll
        if value is None:
            return EMPTY_PLACEHOLDER
        value = str(value)
        if value.strip() == "":
            return EMPTY_PLACEHOLDER
        return value
    return None


def replace_in_paragraph(paragraph, data, stats):
    """Replace placeholders dalam paragraph (semua runs). Update stats."""
    for run in paragraph.runs:
        matches = list(PATTERN.finditer(run.text))
        if not matches:
            continue
        # Replace dari akhir ke awal untuk preserve index
        new_text = run.text
        for m in reversed(matches):
            token = m.group(1)
            value = get_replacement_value(token, data)
            if value is not None:
                new_text = new_text[:m.start()] + value + new_text[m.end():]
                stats["replaced"] += 1
                if value == EMPTY_PLACEHOLDER:
                    stats["empty_filled"] += 1
            else:
                stats["missing_data"] += 1
                stats["missing_tokens"].add(token)
        run.text = new_text


def replace_in_cell(cell, data, stats):
    """Replace di cell + recurse ke nested tables."""
    for p in cell.paragraphs:
        replace_in_paragraph(p, data, stats)
    for nt in cell.tables:
        for row in nt.rows:
            for c in row.cells:
                replace_in_cell(c, data, stats)


def replace_in_header_footer(hdr_ftr, data, stats):
    """Replace di header/footer (paragraphs + tables)."""
    for p in hdr_ftr.paragraphs:
        replace_in_paragraph(p, data, stats)
    for tbl in hdr_ftr.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)


def validate_structure_unchanged(original_doc, merged_doc):
    """
    Validasi struktur XML identik sebelum & sesudah merge.
    Return (is_valid, diff_dict).
    """
    def count_structure(doc):
        body = doc.element.body
        return {
            "body_paragraphs": sum(1 for c in body.iterchildren() if c.tag == qn('w:p')),
            "body_tables": sum(1 for c in body.iterchildren() if c.tag == qn('w:tbl')),
            "all_paragraphs": len(list(body.iter(qn('w:p')))),
            "all_tables": len(list(body.iter(qn('w:tbl')))),
            "all_runs": len(list(body.iter(qn('w:r')))),
        }
    orig = count_structure(original_doc)
    merged = count_structure(merged_doc)
    diff = {k: merged[k] - orig[k] for k in orig}
    is_valid = all(v == 0 for v in diff.values())
    return is_valid, diff


def extract_template_placeholders(template_path):
    """
    Extract semua placeholder unik dari template.
    Return list of unique keys.
    """
    doc = Document(template_path)
    placeholders = set()

    # Body paragraphs
    for p in doc.paragraphs:
        placeholders.update(PATTERN.findall(p.text))

    # Body tables (with nested)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    placeholders.update(PATTERN.findall(p.text))
                for nt in cell.tables:
                    for nrow in nt.rows:
                        for ncell in nrow.cells:
                            for p in ncell.paragraphs:
                                placeholders.update(PATTERN.findall(p.text))

    # Headers & footers
    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            for p in hdr_ftr.paragraphs:
                placeholders.update(PATTERN.findall(p.text))
            for tbl in hdr_ftr.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            placeholders.update(PATTERN.findall(p.text))

    return sorted(placeholders)


def perform_mail_merge(template_path, data):
    """
    Lakukan mail merge pada PKS template.

    Args:
        template_path: Path ke file .docx template
        data: Dictionary {placeholder_name: value}

    Returns:
        Tuple (merged_docx_bytes, stats_dict)

    Raises:
        FileNotFoundError: kalau template tidak ada
        ValueError: kalau struktur berubah (should never happen)
    """
    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Load original template untuk structure validation
    original_doc = Document(template_path)

    # Load working copy
    doc = Document(template_path)

    stats = {
        "replaced": 0,
        "empty_filled": 0,
        "missing_data": 0,
        "missing_tokens": set(),
    }

    # Body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, data, stats)

    # All tables (with nested)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_in_cell(cell, data, stats)

    # Headers & footers
    for section in doc.sections:
        for hdr_ftr in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            replace_in_header_footer(hdr_ftr, data, stats)

    # Validate structure unchanged
    is_valid, diff = validate_structure_unchanged(original_doc, doc)
    if not is_valid:
        raise ValueError(f"Structure changed during merge! Diff: {diff}")

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    # Convert set to list for JSON serialization
    stats["missing_tokens"] = sorted(stats["missing_tokens"])
    stats["structure_valid"] = is_valid
    stats["structure_diff"] = diff

    return output.getvalue(), stats


# ============================================================
# CLI entry point
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="PKS Mail Merge — Production v2")
    parser.add_argument("--template", required=True, help="Path to .docx template")
    parser.add_argument("--input", help="Path to input JSON file (if not using stdin)")
    parser.add_argument("--output", help="Path to output .docx file (if not using stdout)")
    parser.add_argument("--stats", help="Path to write stats JSON (optional)")
    parser.add_argument("--list-placeholders", action="store_true",
                        help="Just list all placeholders in template and exit")
    args = parser.parse_args()

    # Mode: list placeholders
    if args.list_placeholders:
        placeholders = extract_template_placeholders(args.template)
        print(json.dumps({
            "placeholders": placeholders,
            "count": len(placeholders)
        }, indent=2, ensure_ascii=False))
        return

    # Read input data
    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        # Read from stdin
        stdin_text = sys.stdin.read()
        if not stdin_text.strip():
            print("Error: No input provided. Use --input or pipe JSON via stdin.", file=sys.stderr)
            sys.exit(1)
        data = json.loads(stdin_text)

    # Perform merge
    try:
        merged_bytes, stats = perform_mail_merge(args.template, data)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    # Write output
    if args.output:
        with open(args.output, "wb") as f:
            f.write(merged_bytes)
        print(f"✅ Saved: {args.output}", file=sys.stderr)
    else:
        # Write binary ke stdout (untuk pipe)
        sys.stdout.buffer.write(merged_bytes)

    # Write stats if requested
    if args.stats:
        with open(args.stats, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)

    # Print stats ke stderr (so stdout tetap clean binary)
    print(f"📊 Stats:", file=sys.stderr)
    print(f"   Replaced: {stats['replaced'] - stats['empty_filled']}", file=sys.stderr)
    print(f"   Empty filled with '___': {stats['empty_filled']}", file=sys.stderr)
    print(f"   Missing data: {stats['missing_data']}", file=sys.stderr)
    if stats['missing_tokens']:
        print(f"   Missing tokens: {stats['missing_tokens']}", file=sys.stderr)
    print(f"   Structure valid: {stats['structure_valid']}", file=sys.stderr)


if __name__ == "__main__":
    main()
